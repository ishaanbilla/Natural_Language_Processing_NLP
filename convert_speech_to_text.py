from googletrans import Translator
import speech_recognition as sr
import docx
from gtts import gTTS
#Obtain audio from microphone
try:
	r = sr.Recognizer()
	doc = docx.Document()

	with sr.Microphone() as source:
		print('Please wait for 5 seconds. Calibrating microphone.')
		#Listens to the ambient noise for 5 seconds to create a base noise level
		r.adjust_for_ambient_noise(source, duration=5) 
		print('Say Something:')
		audio = r.listen(source)
		print ('Done!')

	text = r.recognize_google(audio, language = 'en-IN')
	#text = r.recognize_google(audio, language = 'bn-BD')
	#text = r.recognize_google(audio, language = 'en-IN')
	#text = r.recognize_google(audio, language = 'ar-AE')

	file=gTTS(text)
	#Saving the audio file of the data user said	
	filename='C:/ISHAAN/Projects/Speech to text in doc W/temp.mp3'
	file.save(filename)

	print (text)

	print (r.recognize_google(audio,language = 'hi-IN'))
	print (r.recognize_google(audio,language = 'ar'))

except sr.UnknownValueError:
        print("Speech Recognition could not understand audio")
translator = Translator()

doc.add_paragraph(translator.translate(text).text)
doc.save('demo_original.docx')

destination_languages = {
	'Arabic' : 'ar'
}


translator = Translator()
'''
Saving the speech data into a Microsoft Word document file with the text in original language
and the language that we want to convert to
'''
for key, value in destination_languages.items():
    doc.add_paragraph(translator.translate(text, dest=value).text)
    doc.save('demo_arabic.docx')

